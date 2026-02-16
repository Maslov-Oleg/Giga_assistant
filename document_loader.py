# document_loader.py
import os
from pathlib import Path

# Для работы с Word документами
try:
    from docx import Document
except ImportError:
    print("Установите python-docx: pip install python-docx")
    raise

def load_document(filepath: str) -> str:
    """
    Загружает документ и возвращает его текст.
    Поддерживает .txt и .docx форматы.
    """
    filepath = Path(filepath)
    
    if not filepath.exists():
        raise FileNotFoundError(f"Файл {filepath} не найден")
    
    # Определяем тип файла по расширению
    extension = filepath.suffix.lower()
    
    if extension == '.txt':
        return _load_txt(filepath)
    elif extension == '.docx':
        return _load_docx(filepath)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {extension}. Используйте .txt или .docx")

def _load_txt(filepath: Path) -> str:
    """Загружает текстовый файл"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def _load_docx(filepath: Path) -> str:
    """Загружает Word документ"""
    doc = Document(filepath)
    # Собираем весь текст из документа
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    # Также можем добавить текст из таблиц, если нужно
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return '\n'.join(full_text)
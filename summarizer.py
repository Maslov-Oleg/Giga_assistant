import os
import sys
import re
import subprocess
import tempfile
from docx2pdf import convert
from docx import Document
from docx.shared import Inches
from gigachat import GigaChat
from gigachat.models import Chat, Messages, MessagesRole


class Summarizer:
    """
    Класс для создания отчёта по конференции с использованием GigaChat
    """

    def __init__(self, api_key: str):
        """
        Инициализация с использование библиотеки gigachat

        Args:
            api_key: Api ключ для доступа к GC
        """
        self.client = GigaChat(credentials=api_key, verify_ssl_certs=False)

    def read_docx(self, file_path: str) -> str:
        """
            Чтение текста из docx файла
        """

        doc = Document(file_path)
        return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

    def merge_texts(self, file_paths: list) -> str:
        """
            Объединение текста из docx файлов
        """

        texts = []
        for path in file_paths:
            text = self.read_docx(path)
            texts.append(f"--- {os.path.basename(path)} ---\n{text}")
        return '\n\n'.join(texts)

    def get_summary_and_code(self, conf_text: str) -> tuple:
        """
            Отправка запроса в GC и получение выжимки с кодом.
        """
        # Системный промпт
        system_content = """Ты — ассистент для глубокого анализа конференций. Твоя задача — обработать два документа: расшифровку выступления спикера и список вопросов-ответов (Q&A), а затем подготовить структурированный отчёт.

**Важно!** В отчёте должно быть три логические части, как описано ниже.

**Часть 1. Краткая выжимка из речи спикера.**
Сделай лаконичное, но содержательное резюме основной части выступления. Выдели ключевые темы, основные мысли и выводы, которые озвучил спикер. Избегай общих фраз, сконцентрируйся на сути.

**Часть 2. Аналитика вопросов участников (самое важное!).**
Не перечисляй все вопросы подряд. Вместо этого:
1. Проанализируй общую массу вопросов из второго файла и выяви основные **тенденции и тематики**. Например: "Больше всего участников волновали вопросы безопасности и юридической ответственности (35%)", "Второй по популярности темой стала интеграция ИИ с устаревшими системами (25%)", "Также прозвучал ряд вопросов о практической пользе для конкретных отраслей (20%)" и т.д.
2. Сгруппируй вопросы по этим тематикам и определи примерное процентное соотношение каждой темы.
3. Для **каждой выделенной тематики** приведи 1-2 наиболее ярких и показательных примера вопросов и соответствующие ответы на них из предоставленного файла.
4. Представь эту часть в виде связного аналитического текста.

**Часть 3. Код для круговой диаграммы.**
Напиши код на Python для создания **строго круговой диаграммы (pie chart)**, которая визуализирует процентное соотношение тематик вопросов, которые ты выявил в Части 2.

Данные для диаграммы:
- Названия секторов: тематики вопросов.
- Размеры секторов: проценты из Части 2, которые ты сам выявил (т.е. если ты нашёл 4 тематики с разделением по процентам для каждого, то ЭТИ же проценты идут и для обозначения этих же тем в диаграмме)

Требования к коду:
- Использовать ТОЛЬКО библиотеку `matplotlib`.
- Сохранять диаграмму в файл с именем `"chart.png"`. **Категорически нельзя использовать `plt.show()`.**
- Сделать диаграмму максимально читаемой:
  * Размер фигуры: `plt.figure(figsize=(10, 8))`
  * Добавить заголовок: `plt.title("Тематика вопросов участников", fontsize=14, pad=20)`
  * Использовать `plt.tight_layout()` перед сохранением
  * В `plt.savefig()` обязательно добавить `bbox_inches='tight', dpi=100`
- Для подписей секторов используй `autopct='%1.1f%%'` для отображения процентов
- Если названия тем длинные, используй легенду (`plt.legend()`) вместо подписей на секторах
- Код должен быть самодостаточным и готовым к выполнению.

**ВАЖНО:**
- Код должен создавать именно КРУГОВУЮ ДИАГРАММУ, а не столбчатую или другую
- Данные для диаграммы (названия секторов и проценты) должны быть явно прописаны в коде на основе твоего анализа из Части 2
- Никаких сложных вычислений или чтения данных из файлов

**Структура финального ответа:**
Сначала идет текст с Частью 1 и Частью 2. После него, отделенный тремя обратными кавычками с пометкой `python`, начинается код из Части 3. Никаких лишних фраз после аналитики и перед кодом быть не должно.
"""

        # Сообщение пользователя
        user_content = f"Текст конференции для анализа:\n\n{conf_text}"

        # запрос с ролями
        messages = [
            Messages(role=MessagesRole.SYSTEM, content=system_content),
            Messages(role=MessagesRole.USER, content=user_content)
        ]

        response = self.client.chat(Chat(messages=messages))

        raw_response = response.choices[0].message.content

        return self._parse_response(raw_response)

    def _parse_response(self, response: str) -> tuple:
        """
        Парсинг ответа от GC - ищем текст и код в ```python
        """

        code_pattern = r'```python\n(.*?)```'
        code_match = re.search(code_pattern, response, re.DOTALL)

        if code_match:
            code = code_match.group(1).strip()
            summary = response[:code_match.start()].strip()
        else:
            code_match = re.search(r'```\n(.*?)```', response, re.DOTALL)
            if code_match:
                code = code_match.group(1).strip()
                summary = response[:code_match.start()].strip()
            else:
                parts = response.split('\n\n')
                if len(parts) > 1 and any('import' in p for p in parts[-1]):
                    summary = '\n\n'.join(parts[:-1])
                    code = parts[-1]
                else:
                    summary = response
                    code = ""

        # очистка summary
        summary = summary.strip()
        summary = re.sub(r'^#+\s*', '', summary, flags=re.MULTILINE)

        # удаление сообщения про код
        patterns_to_remove = [
            r'Часть\s*\d+\.?\s*Задание\s*на\s*генерацию\s*кода\s*для\s*диаграммы\.?\s*\n?',
            r'Часть\s*\d+\.?\s*Код\s*для\s*круговой\s*диаграммы\.?\s*\n?',
            r'Часть\s*\d+\.?\s*Код.*?диаграмм.*?\n',
            r'Задание\s*на\s*генерацию\s*кода\s*для\s*диаграммы\.?\s*\n?',
            r'Код\s*для\s*круговой\с*диаграммы\.?\s*\n?',
            r'^Часть\s*\d+\.?\s*.*?(?:код|диаграмм|chart|code).*?\n',
        ]

        for pattern in patterns_to_remove:
            summary = re.sub(pattern, '', summary, flags=re.IGNORECASE | re.MULTILINE)

        # доп.очистка: убираем лишние пустые строки
        summary = re.sub(r'\n\s*\n\s*\n+', '\n\n', summary)
        summary = summary.strip()

        return summary, code

    def generate_chart(self, code: str) -> bool:
        """
        Выполнение кода и создание диаграммы с использованием текущего интерпретатора Python
        """
        if not code or len(code) < 10:
            return False

        # проверка, нет ли в коде лишних маркеров
        if code.startswith('python'):
            code = code[6:].lstrip()

        # путь к текущему интерпретатору Python
        python_path = sys.executable

        # сохраняем код во временный файл
        with tempfile.NamedTemporaryFile(mode='w', suffix='.py', encoding="utf-8", delete=False) as f:
            code_header = """# -*- coding: utf-8 -*-
import sys
import os
# Убеждаемся, что диаграмма сохраняется
"""

            if 'plt.savefig' not in code:
                if 'plt.show()' in code:
                    code = code.replace('plt.show()', 'plt.savefig("chart.png")')
                else:
                    code = code.rstrip() + '\n\n# Сохраняем диаграмму\nplt.savefig("chart.png", dpi=100, bbox_inches="tight")'


            full_code = code_header + code
            f.write(full_code)
            temp_file = f.name

        try:
            # тот же Python, что и текущий скрипт
            result = subprocess.run(
                [python_path, temp_file],
                capture_output=True,
                text=True,
                timeout=30,
                env=os.environ.copy()  # текущее окружение
            )

            if os.path.exists('chart.png'):
                #size = os.path.getsize('chart.png')
                #(f"✓ Диаграмма создана: {size} байт")
                return True
            else:
                #print("✗ Файл chart.png не создан")
                #if result.stdout:
                #    print(f"Вывод программы:\n{result.stdout}")
                #if result.stderr:
                #    print(f"Ошибки:\n{result.stderr}")
                return False

        except subprocess.TimeoutExpired:
            #print("✗ Превышено время выполнения кода (30 сек)")
            return False
        except Exception as e:
            #print(f"✗ Ошибка при выполнении: {e}")
            return False
        finally:
            # удаление временного файла
            if os.path.exists(temp_file):
                os.unlink(temp_file)

    def create_report(self, docx_files: list, output_file: str = "итоги_конференции.pdf"):
        """
        Основной метод: создаёт полный отчёт с диаграммой формата .pdf .

        Args:
            docx_files: список путей к docx файлам
            output_file: имя итогового файла
        """

        # временный docx файл для конвертации
        temp_docx = output_file.replace('.pdf', '_temp.docx')

        #print("1. Чтение файлов...")
        conference_text = self.merge_texts(docx_files)
        #print(f"    Прочитано {len(conference_text)} символов")

        #print("2. Отправка запроса в Gigachat...")
        summary, code = self.get_summary_and_code(conference_text)
        #print("   ✓ Ответ получен и распарсен")

        #print("3. Создание диаграммы...")
        #print(f"Длина полученного кода: {len(code)} символов")
        #if code:
        #    print("Первые 100 символов кода:")
        #    print(code[:100])
        chart_created = self.generate_chart(code)

        #print("4. Формирование временного документа...")
        doc = Document()

        doc.add_heading('Итоги конференции', 0)

        doc.add_heading('Краткое содержание:', level=1)
        for paragraph in summary.split('\n'):
            if paragraph.strip():
                if paragraph.strip().startswith('**') and paragraph.strip().endswith('**'):
                    doc.add_heading(paragraph.strip('*').strip(), level=2)
                else:
                    doc.add_paragraph(paragraph)

        if chart_created:
            doc.add_heading('Визуализация:', level=1)
            doc.add_picture('chart.png', width=Inches(6))
        else:
            doc.add_heading('Визуализация не создана', level=1)
            doc.add_paragraph('Не удалось сгенерировать диаграмму по данным конференции.')

        doc.save(temp_docx)
        #print(f"4. ✓ Готово! Временный отчёт сохранён в {temp_docx}")

        #print("5. Конвертация в PDF...")
        convert(temp_docx, output_file)
        if os.path.exists(output_file):
            pdf_size = os.path.getsize(output_file)
            #print(f"   ✓ PDF успешно создан! Размер: {pdf_size} байт")

        # ШАГ 6: Очистка временных файлов
        #print("\n6. Очистка временных файлов...")
        #if temp_docx and os.path.exists(temp_docx):
        #    try:
        #        os.unlink(temp_docx)
        #        print("   ✓ Временный DOCX удалён")
        #    except Exception as e:
        #        print(f"   ⚠ Не удалось удалить временный файл: {e}")
#
        #    # Удаляем диаграмму, если она была создана
        #if chart_created and os.path.exists('chart.png'):
        #    try:
        #        os.unlink('chart.png')
        #        print("   ✓ Временная диаграмма удалена")
        #    except Exception as e:
        #        print(f"   ⚠ Не удалось удалить диаграмму: {e}")